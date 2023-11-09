from django.test import TestCase

# Create your tests here.
import docx

doc = docx.Document("/Users/donaghymo/Desktop/成语/成语  文字/成语4-6级.docx")

for i in range(len(doc.paragraphs)):
    print("----这是第{}段------------".format(i))
    print(doc.paragraphs[i].text)

def transfer_doc2string(file_location):
    """
    将doc转换成string
    """
    doc = docx.Document(file_location)
    lines = ""
    for i in range(len(doc.paragraphs)):
        lines=lines+doc.paragraphs[i].text
    return lines

if __name__ == "__main__":
    print(transfer_doc2string("/Users/donaghymo/Desktop/成语/成语  文字/成语4-6级.docx"))